"""YAML → DOCX assembly."""
from __future__ import annotations

from pathlib import Path

import yaml
from docx import Document
from docx.shared import Mm
from PIL import Image, ImageDraw, ImageFont


def _style_ok(doc: Document, name: str | None) -> str | None:
    if not name:
        return None
    try:
        doc.styles[name]
        return name
    except KeyError:
        return None


def _ensure_img(img_path: Path, label: str) -> Path:
    if img_path.exists():
        return img_path
    if img_path.suffix.lower() not in (
        ".png", ".jpg", ".jpeg", ".bmp", ".gif",
    ):
        img_path = img_path.with_suffix(".png")
    img_path.parent.mkdir(parents=True, exist_ok=True)
    im = Image.new("RGB", (128, 64), (245, 245, 245))
    d = ImageDraw.Draw(im)
    d.rectangle([(0, 0), (127, 63)], outline=(80, 80, 80), width=2)
    try:
        fnt = ImageFont.truetype("arial.ttf", 12)
    except Exception:
        fnt = ImageFont.load_default()
    d.text((10, 20), Path(label).name[:20], fill=(20, 20, 20), font=fnt)
    im.save(img_path, "PNG")
    return img_path


def build_docx_from_yaml(
    yaml_path: str, template_docx: str, out_docx: str
) -> None:
    y = yaml.safe_load(Path(yaml_path).read_text(encoding="utf-8"))
    doc = Document(template_docx)
    base = Path(yaml_path).resolve().parent

    for blk in y.get("blocks", []):
        t = blk.get("type")

        if t == "heading":
            p = doc.add_paragraph(blk.get("text", ""))
            s = _style_ok(doc, blk.get("style"))
            if not s:
                lv = blk.get("level")
                if isinstance(lv, int) and lv > 0:
                    s = _style_ok(doc, f"Heading {lv}")
            if s:
                p.style = s

        elif t == "paragraph":
            p = doc.add_paragraph(blk.get("text", ""))
            s = _style_ok(doc, blk.get("style"))
            if s:
                p.style = s

        elif t == "image":
            src = blk.get("src")
            if not src:
                continue

            src_str = str(src).lstrip("/\\")
            src_path = Path(src_str)

            if src_path.is_absolute():
                ip = _ensure_img(src_path, str(src))
            else:
                ip = _ensure_img(base / src_path, str(src))

            p = doc.add_paragraph()
            s = _style_ok(doc, blk.get("style"))
            if s:
                p.style = s
            run = p.add_run()
            w = blk.get("width_mm")
            if w:
                run.add_picture(str(ip), width=Mm(w))
            else:
                run.add_picture(str(ip))
            cap = blk.get("caption")
            if isinstance(cap, dict) and cap.get("text"):
                cp = doc.add_paragraph(cap["text"])
                cs = _style_ok(doc, cap.get("style"))
                if cs:
                    cp.style = cs

        elif t == "table":
            rows = blk.get("rows") or []
            if not rows:
                continue
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    tbl.cell(i, j).text = "" if val is None else str(val)
            ts = blk.get("style")
            if ts:
                try:
                    doc.styles[ts]
                    tbl.style = ts
                except Exception:
                    pass

    doc.save(out_docx)